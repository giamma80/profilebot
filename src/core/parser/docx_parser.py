"""DOCX parser for ProfileBot CV ingestion."""

from __future__ import annotations

import logging
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from src.core.parser.metadata_extractor import extract_metadata
from src.core.parser.schemas import CVMetadata, ExperienceItem, ParsedCV, SkillSection
from src.core.parser.section_detector import detect_sections

logger = logging.getLogger(__name__)


class CVParseError(Exception):
    """Raised when a CV cannot be parsed."""


@dataclass(frozen=True)
class ParsedSections:
    """Structured sections extracted from the CV."""

    skills: list[str]
    experience: list[str]
    education: list[str]
    certifications: list[str]
    raw_text: str


class DocxParser:
    """Parser for DOCX CV files."""

    def parse(self, file_path: str | Path) -> ParsedCV:
        """Parse a DOCX CV and return a structured ParsedCV object."""
        path = Path(file_path)
        if not path.exists():
            raise CVParseError(f"File not found: {path}")

        start_time = time.perf_counter()

        try:
            document = Document(str(path))
        except PackageNotFoundError as exc:
            raise CVParseError(f"Invalid or corrupted DOCX file: {path}") from exc
        except Exception as exc:  # pragma: no cover - defensive
            raise CVParseError(f"Failed to read DOCX file: {path}") from exc

        lines = list(self._extract_lines(document))
        raw_text = "\n".join(lines).strip()

        if not raw_text:
            metadata = self._build_metadata(path, raw_text)
            parsed = ParsedCV(
                metadata=metadata,
                skills=None,
                experiences=[],
                education=[],
                certifications=[],
                raw_text="",
            )
            self._log_parse_result(parsed, ParsedSections([], [], [], [], raw_text), start_time)
            return parsed

        sections = self._extract_sections(lines, raw_text)
        metadata = self._build_metadata(path, raw_text)

        parsed = ParsedCV(
            metadata=metadata,
            skills=self._parse_skills(sections.skills),
            experiences=self._parse_experiences(sections.experience),
            education=sections.education,
            certifications=sections.certifications,
            raw_text=sections.raw_text,
        )
        self._log_parse_result(parsed, sections, start_time)
        return parsed

    def _extract_lines(self, document: Document) -> Iterable[str]:
        """Extract text lines from paragraphs and tables in a DOCX document."""
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name and style_name.lower().startswith("heading"):
                yield text
                continue
            yield text

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        for line in text.splitlines():
                            line = line.strip()
                            if line:
                                yield line

    def _extract_sections(self, lines: list[str], raw_text: str) -> ParsedSections:
        """Detect and group content into sections."""
        detected = detect_sections(lines)

        return ParsedSections(
            skills=detected.get("skills", []),
            experience=detected.get("experience", []),
            education=detected.get("education", []),
            certifications=detected.get("certifications", []),
            raw_text=raw_text,
        )

    def _log_parse_result(
        self, parsed: ParsedCV, sections: ParsedSections, start_time: float
    ) -> None:
        parse_time_ms = int((time.perf_counter() - start_time) * 1000)
        sections_found = {
            "skills": len(sections.skills),
            "experience": len(sections.experience),
            "education": len(sections.education),
            "certifications": len(sections.certifications),
        }
        logger.info(
            "cv_parsed",
            extra={
                "cv_id": parsed.metadata.cv_id,
                "file_name": parsed.metadata.file_name,
                "sections_found": sections_found,
                "parse_time_ms": parse_time_ms,
            },
        )

    def _build_metadata(self, path: Path, raw_text: str) -> CVMetadata:
        """Extract metadata from the document content."""
        metadata = extract_metadata(raw_text)
        return CVMetadata(
            cv_id=metadata.cv_id,
            file_name=path.name,
            full_name=metadata.full_name,
            current_role=metadata.current_role,
            parsed_at=metadata.parsed_at,
        )

    def _parse_skills(self, lines: list[str]) -> SkillSection | None:
        """Parse the skills section into raw text and keywords."""
        raw = "\n".join(lines).strip()
        if not raw:
            return None
        keywords = self._split_keywords(raw)
        return SkillSection(raw_text=raw, skill_keywords=keywords)

    def _parse_experiences(self, lines: list[str]) -> list[ExperienceItem]:
        """Parse the experience section into structured items."""
        if not lines:
            return []

        items: list[ExperienceItem] = []
        buffer: list[str] = []

        for line in lines:
            if self._is_new_experience_line(line) and buffer:
                items.append(self._buffer_to_experience(buffer))
                buffer = [line]
            else:
                buffer.append(line)

        if buffer:
            items.append(self._buffer_to_experience(buffer))

        return items

    def _buffer_to_experience(self, lines: list[str]) -> ExperienceItem:
        """Convert a list of lines into an ExperienceItem."""
        description = "\n".join(lines).strip()
        lowered = description.lower()
        is_current = any(token in lowered for token in ("present", "current", "oggi", "attuale"))
        return ExperienceItem(
            company=None,
            role=None,
            start_date=None,
            end_date=None,
            description=description,
            is_current=is_current,
        )

    def _is_new_experience_line(self, line: str) -> bool:
        """Heuristic to detect a new experience entry."""
        return line.isupper() or "–" in line or "-" in line

    def _split_keywords(self, text: str) -> list[str]:
        """Split a skills string into keywords."""
        tokens = []
        for chunk in text.replace(";", ",").replace("|", ",").split(","):
            cleaned = chunk.strip()
            if cleaned:
                tokens.append(cleaned)
        return tokens


def parse_docx(file_path: str | Path) -> ParsedCV:
    """Convenience function to parse a DOCX CV file."""
    return DocxParser().parse(file_path)


__all__ = ["CVParseError", "DocxParser", "parse_docx"]
