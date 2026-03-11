"""Tests for the DOCX CV parser (US-003)."""

from __future__ import annotations

import re
import shutil
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from src.core.parser.docx_parser import CVParseError, parse_docx, parse_docx_bytes
from src.core.parser.schemas import ParsedCV

FIXTURES_DIR = Path(__file__).parent / "fixtures" / "sample_cvs"
CAMPIONE_DIR = FIXTURES_DIR / "campione"


def _list_docx_fixtures() -> list[Path]:
    return sorted(FIXTURES_DIR.glob("*.docx"))


def _copy_fixture_with_res_id(tmp_path: Path, fixture_path: Path, res_id: int) -> Path:
    destination = tmp_path / f"{res_id}_{fixture_path.name}"
    shutil.copy(fixture_path, destination)
    return destination


def _parse_campione_docx(file_name: str) -> ParsedCV:
    path = CAMPIONE_DIR / file_name
    if not path.exists():
        pytest.skip(f"Campione fixture not available: {path}")
    match = re.search(r"\d+", file_name)
    if not match:
        raise ValueError(f"Missing res_id in filename: {file_name}")
    res_id = int(match.group(0))
    return parse_docx_bytes(path.read_bytes(), res_id)


def test_fixtures_exist() -> None:
    """Ensure we have enough sample CVs to validate parsing."""
    fixtures = _list_docx_fixtures()
    assert len(fixtures) >= 5, "Expected at least 5 DOCX fixtures in tests/fixtures/sample_cvs"


@pytest.mark.parametrize("docx_path", _list_docx_fixtures())
def test_parse_docx_returns_parsed_cv(docx_path: Path, tmp_path: Path) -> None:
    """Parse each DOCX fixture and validate basic structure."""
    res_id = 12345
    res_id_path = _copy_fixture_with_res_id(tmp_path, docx_path, res_id)
    parsed = parse_docx(res_id_path)

    assert parsed is not None
    assert parsed.metadata is not None
    assert parsed.metadata.cv_id
    assert parsed.metadata.file_name
    assert parsed.metadata.res_id == res_id

    if parsed.skills is not None:
        assert isinstance(parsed.skills.raw_text, str)
        assert isinstance(parsed.skills.skill_keywords, list)

    assert isinstance(parsed.experiences, list)
    assert isinstance(parsed.education, list)
    assert isinstance(parsed.certifications, list)
    assert isinstance(parsed.raw_text, str)
    assert parsed.raw_text.strip() != ""


def test_parse_invalid_docx_raises(tmp_path: Path) -> None:
    """Invalid DOCX files should raise a CVParseError."""
    invalid_path = tmp_path / "invalid.docx"
    invalid_path.write_bytes(b"not a docx file")

    with pytest.raises(CVParseError):
        parse_docx(invalid_path)


def test_parse_docx_bytes__valid_docx__returns_parsed_cv() -> None:
    document = Document()
    document.add_paragraph("Test CV")
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_docx_bytes(buffer.getvalue(), 12345)

    assert parsed.metadata.res_id == 12345
    assert parsed.metadata.file_name == "12345_unknown.docx"
    assert "Test CV" in parsed.raw_text


def test_parse_docx_bytes__invalid_docx__raises_parse_error() -> None:
    with pytest.raises(CVParseError):
        parse_docx_bytes(b"not a docx file", 12345)


def test_parse_docx__filename_with_res_id__sets_metadata_res_id(tmp_path: Path) -> None:
    """Valid filename prefixes should populate res_id."""
    docx_path = tmp_path / "12345_mario_rossi.docx"
    document = Document()
    document.add_paragraph("Test CV")
    document.save(docx_path)

    parsed = parse_docx(docx_path)

    assert parsed.metadata.res_id == 12345


def test_parse_docx__missing_res_id__raises_parse_error(tmp_path: Path) -> None:
    """Missing res_id prefix should raise CVParseError."""
    docx_path = tmp_path / "mario_rossi.docx"
    document = Document()
    document.add_paragraph("Test CV")
    document.save(docx_path)

    with pytest.raises(CVParseError):
        parse_docx(docx_path)


@pytest.mark.parametrize(
    "filename,expected", [("12345_mario_rossi.docx", 12345), ("99999_a_b.docx", 99999)]
)
def test_parse_docx__res_id_edge_cases__parses_numeric_prefix(
    tmp_path: Path, filename: str, expected: int
) -> None:
    """Numeric prefixes should be parsed even with short names."""
    docx_path = tmp_path / filename
    document = Document()
    document.add_paragraph("Test CV")
    document.save(docx_path)

    parsed = parse_docx(docx_path)

    assert parsed.metadata.res_id == expected


def test_parse_docx__leading_zeros__strips_to_int(tmp_path: Path) -> None:
    """Leading zeros in res_id should be parsed as int."""
    docx_path = tmp_path / "000123_mario_rossi.docx"
    document = Document()
    document.add_paragraph("Test CV")
    document.save(docx_path)

    parsed = parse_docx(docx_path)

    assert parsed.metadata.res_id == 123


def test_parse_standard_cv_has_sections(tmp_path: Path) -> None:
    docx_path = _copy_fixture_with_res_id(tmp_path, FIXTURES_DIR / "cv_standard.docx", 12345)
    parsed = parse_docx(docx_path)
    assert parsed.skills is not None
    assert parsed.experiences
    assert parsed.education
    assert parsed.certifications


def test_parse_cv_with_tables_includes_skills(tmp_path: Path) -> None:
    docx_path = _copy_fixture_with_res_id(tmp_path, FIXTURES_DIR / "cv_with_tables.docx", 12345)
    parsed = parse_docx(docx_path)
    assert parsed.skills is not None
    assert parsed.skills.skill_keywords


def test_parse_unstructured_cv_has_raw_text(tmp_path: Path) -> None:
    docx_path = _copy_fixture_with_res_id(tmp_path, FIXTURES_DIR / "cv_unstructured.docx", 12345)
    parsed = parse_docx(docx_path)
    assert parsed.raw_text.strip() != ""


def test_parse_italian_chars_cv(tmp_path: Path) -> None:
    docx_path = _copy_fixture_with_res_id(tmp_path, FIXTURES_DIR / "cv_italian_chars.docx", 12345)
    parsed = parse_docx(docx_path)
    assert "è" in parsed.raw_text or "à" in parsed.raw_text or "ù" in parsed.raw_text


def test_parse_minimal_cv_skills(tmp_path: Path) -> None:
    docx_path = _copy_fixture_with_res_id(tmp_path, FIXTURES_DIR / "cv_minimal.docx", 12345)
    parsed = parse_docx(docx_path)
    assert parsed.skills is not None
    assert parsed.skills.skill_keywords


def test_parse_cv__multipart_name__extracts_full_name() -> None:
    docx_path = FIXTURES_DIR / "curriculum_212255.docx"
    if not docx_path.exists():
        pytest.skip(f"Fixture not available: {docx_path}")
    parsed = parse_docx_bytes(docx_path.read_bytes(), 212255, filename=docx_path.name)
    assert parsed.metadata.full_name == "Giuseppe Alessandro DE BLASIO"


def test_parse_campione_cv__uppercase_surname__extracts_name_and_role() -> None:
    parsed = _parse_campione_docx("curriculum_12369.docx")
    assert parsed.metadata.full_name == "Andrea AMOROSO"
    assert parsed.metadata.current_role == "CAD ENGINEER - CATIA"


def test_parse_campione_cv__uppercase_role_line__extracts_role() -> None:
    parsed = _parse_campione_docx("curriculum_218087.docx")
    assert parsed.metadata.full_name == "Luca GANDOLFI"
    assert parsed.metadata.current_role == "SOFTWARE ENGINEER - C, C++ (OBJECT ORIENTED)"
